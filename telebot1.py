from random import randint
import telebot
from docx import Document
global mas
mas=None
mas=[]
bot = telebot.TeleBot('1055740850:AAE6PrRsj4pRefeGnRX_UlXMTtvKYTHNxpw')

@bot.message_handler(content_types=['text'])
def qq(message):
    v=randint(150,199)
    global mas
    wordDoc = Document('us.docx')
    k=set()
    h=0
    if (message.text).lower()=='vv3':
     for i in range(1,8):
         if wordDoc.tables[v].rows[i].cells[0].text == '1':
                  k.add('f'+wordDoc.tables[v].rows[i].cells[-1].text)

             
         else:
             k.add(wordDoc.tables[v].rows[i].cells[1].text)
     bot.send_message(message.from_user.id,wordDoc.tables[v].rows[0].cells[-1].text)
     for j in k:
        h+=1
        if j[0]=='f':
             bot.send_message(message.from_user.id, f'ответ {h} : {j[1:len(j)]}')
             mas.append(h)     
        else:
               bot.send_message(message.from_user.id, f'ответ {h} : {j}')
     h=0
     bot.send_message(message.from_user.id,'Напишите 3 цифры ответа через пробел ')
     bot.register_next_step_handler(message, writetask)

def writetask(message):
    global mas
    k=message.text
    print(k)
    n=[int(x) for x in k.split()]


    if n[0] in mas and n[1] in mas and n[2] in mas :
        bot.send_message(message.from_user.id,'Правильно ')
        print(mas)
    else:
        bot.send_message(message.from_user.id,'Не правильно ')
        bot.send_message(message.from_user.id, f'правильно: {mas} ')
        print(mas)
    mas=[]
        
bot.polling(none_stop=True, interval=0)   
    
        
        
    
    
